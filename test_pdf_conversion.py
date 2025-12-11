"""Test portable LibreOffice PDF conversion."""
import sys
import os
from pathlib import Path

# Add project to path
sys.path.insert(0, r'c:\Users\pc\autoarendt')

from services.format_converter import FormatConverter

# Create a simple test Word document
test_dir = Path(r'c:\Users\pc\autoarendt\test_libreoffice')
test_dir.mkdir(exist_ok=True)

# Create a simple docx file for testing
try:
    from docx import Document
    
    doc = Document()
    doc.add_heading('LibreOffice Test', 0)
    doc.add_paragraph('This is a test document to verify LibreOffice PDF conversion.')
    doc.add_paragraph('If you can read this as a PDF, the portable LibreOffice is working correctly!')
    
    test_docx = test_dir / 'test.docx'
    doc.save(str(test_docx))
    print(f'[OK] Created test document: {test_docx}')
    
    # Test conversion
    converter = FormatConverter()
    output_pdf = converter.convert(
        str(test_docx),
        'pdf',
        str(test_dir)
    )
    
    if os.path.exists(output_pdf):
        print(f'[OK] PDF conversion successful: {output_pdf}')
        print(f'   PDF size: {os.path.getsize(output_pdf)} bytes')
    else:
        print(f'[ERROR] PDF not created')
        
except Exception as e:
    print(f'[ERROR] Error: {e}')
    import traceback
    traceback.print_exc()
