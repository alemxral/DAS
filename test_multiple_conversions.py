"""Test multiple conversions to see true speed."""
import sys
import os
import time
from pathlib import Path

sys.path.insert(0, r'c:\Users\pc\autoarendt')

from services.format_converter import FormatConverter
from docx import Document

test_dir = Path(r'c:\Users\pc\autoarendt\test_libreoffice')
test_dir.mkdir(exist_ok=True)

converter = FormatConverter()

times = []
for i in range(3):
    # Create a test document
    doc = Document()
    doc.add_heading(f'Test Document {i+1}', 0)
    doc.add_paragraph('Testing conversion speed.')
    for j in range(3):
        doc.add_paragraph(f'Content paragraph {j+1}')
    
    test_docx = test_dir / f'test_{i+1}.docx'
    doc.save(str(test_docx))
    
    # Convert
    print(f"\nConversion {i+1}/3...")
    start = time.time()
    try:
        pdf_path = converter.convert(str(test_docx), 'pdf', str(test_dir))
        elapsed = time.time() - start
        times.append(elapsed)
        print(f"✅ Completed in {elapsed:.2f}s")
    except Exception as e:
        print(f"❌ Failed: {e}")

if times:
    print("\n" + "=" * 60)
    print("Results:")
    print("=" * 60)
    for i, t in enumerate(times, 1):
        print(f"Conversion {i}: {t:.2f}s")
    print(f"\nAverage: {sum(times)/len(times):.2f}s")
    print(f"First run: {times[0]:.2f}s (includes initialization)")
    if len(times) > 1:
        avg_subsequent = sum(times[1:]) / len(times[1:])
        print(f"Subsequent runs: {avg_subsequent:.2f}s average")
