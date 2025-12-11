"""
Template Processor Service
Processes templates in Word, Excel, and .msg formats with placeholder substitution.
"""
import os
import re
from pathlib import Path
from typing import Dict, List
import shutil

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    Document = None

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter, range_boundaries
except ImportError:
    openpyxl = None

try:
    import win32com.client
    import pythoncom
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    pythoncom = None


class TemplateProcessor:
    """Processes templates with placeholder substitution."""
    
    VARIABLE_PATTERN = re.compile(r'##([^#]+)##')
    
    def __init__(self):
        """Initialize TemplateProcessor."""
        self.supported_formats = ['.docx', '.xlsx', '.msg']
        # Template caches for performance - avoids reloading same template
        self._docx_cache = {}
        self._xlsx_cache = {}
        print("[TemplateProcessor] Initialized with template caching enabled")
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats
    
    def extract_template_variables(self, template_path: str) -> List[str]:
        """
        Extract all ##variable## placeholders from a template.
        
        Args:
            template_path: Path to template file
            
        Returns:
            List of unique variable names found in template
        """
        ext = Path(template_path).suffix.lower()
        
        if ext == '.docx':
            return self._extract_variables_from_docx(template_path)
        elif ext == '.xlsx':
            return self._extract_variables_from_xlsx(template_path)
        elif ext == '.msg':
            return self._extract_variables_from_msg(template_path)
        else:
            raise ValueError(f"Unsupported template format: {ext}")
    
    def _extract_variables_from_docx(self, file_path: str) -> List[str]:
        """Extract variables from Word document."""
        if Document is None:
            raise ImportError("python-docx is required for Word templates")
        
        doc = Document(file_path)
        variables = set()
        
        # Check paragraphs
        for para in doc.paragraphs:
            variables.update(self.VARIABLE_PATTERN.findall(para.text))
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    variables.update(self.VARIABLE_PATTERN.findall(cell.text))
        
        # Check headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                variables.update(self.VARIABLE_PATTERN.findall(para.text))
            for para in section.footer.paragraphs:
                variables.update(self.VARIABLE_PATTERN.findall(para.text))
        
        return sorted(list(variables))
    
    def _extract_variables_from_xlsx(self, file_path: str) -> List[str]:
        """Extract variables from Excel workbook."""
        if openpyxl is None:
            raise ImportError("openpyxl is required for Excel templates")
        
        wb = openpyxl.load_workbook(file_path)
        variables = set()
        
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        variables.update(self.VARIABLE_PATTERN.findall(cell.value))
        
        return sorted(list(variables))
    
    def _extract_variables_from_msg(self, file_path: str) -> List[str]:
        """Extract variables from .msg file."""
        if not WIN32_AVAILABLE:
            raise ImportError("pywin32 is required for .msg templates")
        
        variables = set()
        
        # Initialize COM for this thread
        pythoncom.CoInitialize()
        try:
            outlook = win32com.client.Dispatch("Outlook.Application")
            msg = outlook.CreateItemFromTemplate(file_path)
            
            # Check subject
            if msg.Subject:
                variables.update(self.VARIABLE_PATTERN.findall(msg.Subject))
            
            # Check body
            if msg.Body:
                variables.update(self.VARIABLE_PATTERN.findall(msg.Body))
            
            # Check HTML body
            if msg.HTMLBody:
                variables.update(self.VARIABLE_PATTERN.findall(msg.HTMLBody))
            
        except Exception as e:
            raise ValueError(f"Error reading .msg template: {str(e)}")
        finally:
            pythoncom.CoUninitialize()
        
        return sorted(list(variables))
    
    def process_template(self, template_path: str, data: Dict, output_path: str, sheet_name: str = None, auto_adjust_options: Dict = None) -> str:
        """
        Process a template with data substitution.
        
        Args:
            template_path: Path to template file
            data: Dictionary mapping variable names to values
            output_path: Path for output file
            sheet_name: Optional sheet name for Excel templates
            auto_adjust_options: Optional dict with Excel auto-adjust settings
            
        Returns:
            Path to the generated file
        """
        ext = Path(template_path).suffix.lower()
        
        # Debug logging
        print(f"[TemplateProcessor] Processing template: {Path(template_path).name}")
        print(f"[TemplateProcessor] Data keys: {list(data.keys())}")
        print(f"[TemplateProcessor] Data values count: {len(data)}")
        
        if not data:
            print("[TemplateProcessor] WARNING: Empty data dictionary provided!")
        
        if ext == '.docx':
            return self._process_docx_template(template_path, data, output_path)
        elif ext == '.xlsx':
            return self._process_xlsx_template(template_path, data, output_path, sheet_name, auto_adjust_options)
        elif ext == '.msg':
            return self._process_msg_template(template_path, data, output_path)
        else:
            raise ValueError(f"Unsupported template format: {ext}")
    
    def _replace_text_in_paragraph(self, paragraph, data: Dict):
        """Replace variables in a paragraph while preserving formatting."""
        for var_name, value in data.items():
            placeholder = f"##{var_name}##"
            if placeholder in paragraph.text:
                # Replace inline while preserving runs
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    def _process_docx_template(self, template_path: str, data: Dict, output_path: str) -> str:
        """Process Word template with caching for performance."""
        if Document is None:
            raise ImportError("python-docx is required for Word templates")
        
        # Load from cache or disk (5-10ms vs 50-200ms per load)
        if template_path not in self._docx_cache:
            print(f"[TemplateProcessor] Caching Word template: {Path(template_path).name}")
            self._docx_cache[template_path] = Document(template_path)
        
        # Deep copy cached template (fast memory operation ~5ms)
        from copy import deepcopy
        doc = deepcopy(self._docx_cache[template_path])
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            self._replace_text_in_paragraph(para, data)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_text_in_paragraph(para, data)
        
        # Replace in headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                self._replace_text_in_paragraph(para, data)
            for para in section.footer.paragraphs:
                self._replace_text_in_paragraph(para, data)
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path
    
    def _process_xlsx_template(self, template_path: str, data: Dict, output_path: str, sheet_name: str = None, auto_adjust_options: Dict = None) -> str:
        """Process Excel template with caching for performance."""
        if openpyxl is None:
            raise ImportError("openpyxl is required for Excel templates")
        
        # Load from cache or disk
        if template_path not in self._xlsx_cache:
            print(f"[TemplateProcessor] Caching Excel template: {Path(template_path).name}")
            self._xlsx_cache[template_path] = openpyxl.load_workbook(template_path)
        
        # Load a fresh copy from the cached path (avoid deepcopy issues with openpyxl)
        wb = openpyxl.load_workbook(template_path)
        
        try:
            # Track modified cells for auto-adjust
            modified_cells = set()
            
            # Determine which sheets to process
            sheets_to_process = []
            if sheet_name:
                # Process only the specified sheet
                if sheet_name in wb.sheetnames:
                    sheets_to_process = [wb[sheet_name]]
                else:
                    raise ValueError(f"Sheet '{sheet_name}' not found in template")
            else:
                # Process all sheets
                sheets_to_process = wb.worksheets
            
            for sheet in sheets_to_process:
                replacements_made = 0
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            original_value = cell.value
                            new_value = cell.value
                            for var_name, value in data.items():
                                placeholder = f"##{var_name}##"
                                if placeholder in new_value:
                                    new_value = new_value.replace(placeholder, str(value))
                                    replacements_made += 1
                            if new_value != cell.value:
                                cell.value = new_value
                                # Track modified cell (sheet_title, row, col) tuple
                                modified_cells.add((sheet.title, cell.row, cell.column))
                
                print(f"[TemplateProcessor] Sheet '{sheet.title}': Made {replacements_made} replacements")
            
            print(f"[TemplateProcessor] Total modified cells: {len(modified_cells)}")
            
            # Apply auto-adjust if options provided
            if auto_adjust_options:
                self._apply_excel_auto_adjust(wb, sheets_to_process, modified_cells, auto_adjust_options)
            
            # Save the workbook
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wb.save(output_path)
            return output_path
        finally:
            # Always close the workbook to release file handle
            wb.close()
    
    def _apply_excel_auto_adjust(self, wb, sheets: list, modified_cells: set, options: Dict):
        """
        Apply auto-adjust to Excel cells based on options.
        
        Args:
            wb: Workbook instance
            sheets: List of sheets to process
            modified_cells: Set of (sheet_title, row, col) tuples for modified cells
            options: Dict with auto_adjust_height, auto_adjust_width, adjust_range
        """
        auto_adjust_height = options.get('auto_adjust_height', False)
        auto_adjust_width = options.get('auto_adjust_width', False)
        adjust_range = options.get('adjust_range', None)
        
        for sheet in sheets:
            # Get cells to adjust
            cells_to_adjust = set()
            
            if adjust_range:
                # Parse range like "A1:E10"
                try:
                    min_col, min_row, max_col, max_row = range_boundaries(adjust_range)
                    for row in range(min_row, max_row + 1):
                        for col in range(min_col, max_col + 1):
                            cells_to_adjust.add((row, col))
                    print(f"[TemplateProcessor] Auto-adjusting range {adjust_range} in sheet '{sheet.title}'")
                except Exception as e:
                    print(f"[TemplateProcessor] Warning: Invalid adjust_range '{adjust_range}': {e}")
                    continue
            else:
                # Use modified cells for this sheet
                for sheet_title, row, col in modified_cells:
                    if sheet_title == sheet.title:
                        cells_to_adjust.add((row, col))
                if cells_to_adjust:
                    print(f"[TemplateProcessor] Auto-adjusting {len(cells_to_adjust)} modified cells in sheet '{sheet.title}'")
            
            if not cells_to_adjust:
                continue
            
            # Auto-adjust heights
            if auto_adjust_height:
                rows_to_adjust = set(row for row, col in cells_to_adjust)
                for row_num in rows_to_adjust:
                    # Ensure RowDimension exists before accessing
                    if row_num not in sheet.row_dimensions:
                        from openpyxl.worksheet.dimensions import RowDimension
                        sheet.row_dimensions[row_num] = RowDimension(sheet, index=row_num)
                    sheet.row_dimensions[row_num].height = None  # Auto height
            
            # Auto-adjust widths
            if auto_adjust_width:
                cols_to_adjust = set(col for row, col in cells_to_adjust)
                for col_num in cols_to_adjust:
                    # Calculate optimal width based on cell content
                    max_length = 0
                    column_letter = get_column_letter(col_num)
                    
                    for row_num, cell_col in cells_to_adjust:
                        if cell_col == col_num:
                            cell = sheet.cell(row=row_num, column=col_num)
                            if cell.value:
                                cell_length = len(str(cell.value))
                                max_length = max(max_length, cell_length)
                    
                    # Set column width with some padding (Excel uses character units)
                    adjusted_width = min(max_length + 2, 50)
                    if adjusted_width > 0:
                        sheet.column_dimensions[column_letter].width = adjusted_width
    
    def _process_msg_template(self, template_path: str, data: Dict, output_path: str) -> str:
        """Process .msg template."""
        if not WIN32_AVAILABLE:
            raise ImportError("pywin32 is required for .msg templates")
        
        # Initialize COM for this thread
        pythoncom.CoInitialize()
        try:
            outlook = win32com.client.Dispatch("Outlook.Application")
            msg = outlook.CreateItemFromTemplate(template_path)
            
            # Replace in subject
            if msg.Subject:
                for var_name, value in data.items():
                    placeholder = f"##{var_name}##"
                    msg.Subject = msg.Subject.replace(placeholder, str(value))
            
            # Replace in body
            if msg.Body:
                for var_name, value in data.items():
                    placeholder = f"##{var_name}##"
                    msg.Body = msg.Body.replace(placeholder, str(value))
            
            # Replace in HTML body
            if msg.HTMLBody:
                for var_name, value in data.items():
                    placeholder = f"#{var_name}##"
                    msg.HTMLBody = msg.HTMLBody.replace(placeholder, str(value))
            
            # Save the message
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            msg.SaveAs(output_path)
            
        except Exception as e:
            raise ValueError(f"Error processing .msg template: {str(e)}")
        finally:
            pythoncom.CoUninitialize()
        
        return output_path
    
    def validate_template_data(self, template_path: str, data: Dict) -> Dict:
        """
        Validate that all template variables have corresponding data.
        
        Args:
            template_path: Path to template file
            data: Data dictionary
            
        Returns:
            Validation result with missing/extra variables
        """
        template_vars = set(self.extract_template_variables(template_path))
        data_vars = set(data.keys())
        
        missing_vars = template_vars - data_vars
        extra_vars = data_vars - template_vars
        
        return {
            'is_valid': len(missing_vars) == 0,
            'missing_variables': sorted(list(missing_vars)),
            'extra_variables': sorted(list(extra_vars)),
            'template_variables': sorted(list(template_vars)),
            'data_variables': sorted(list(data_vars))
        }
