"""
Word Operations Service
Handles Word document splitting and merging operations.
"""
import os
from typing import List, Optional
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class WordSplitter:
    """Handles Word document splitting operations."""
    
    def __init__(self, input_docx_path: str):
        """
        Initialize Word splitter.
        
        Args:
            input_docx_path: Path to input Word document
        """
        self.input_docx_path = input_docx_path
        self.doc = Document(input_docx_path)
    
    def split_by_sections(self, output_dir: str, base_name: str = "section") -> List[str]:
        """
        Split Word document by section breaks.
        
        Args:
            output_dir: Directory to save split files
            base_name: Base name for output files
            
        Returns:
            List of paths to created split files
        """
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Get sections
        sections = self.doc.sections
        num_sections = len(sections)
        
        if num_sections <= 1:
            print("Warning: Document has only one section. Cannot split by sections.")
            # Save as single file
            output_file = os.path.join(output_dir, f"{base_name}_1.docx")
            self.doc.save(output_file)
            return [output_file]
        
        output_files = []
        
        # Note: python-docx doesn't support splitting by sections directly
        # This is a limitation - we'll need to convert to PDF first
        print(f"Warning: Word splitting by sections is limited. Document has {num_sections} sections.")
        print("Consider converting to PDF first for better split support.")
        
        # For now, save as single file
        output_file = os.path.join(output_dir, f"{base_name}_complete.docx")
        self.doc.save(output_file)
        output_files.append(output_file)
        
        return output_files
    
    def split_by_pages(self, pages_per_split: int, output_dir: str, base_name: str = "split") -> List[str]:
        """
        Split Word document by page count.
        Note: python-docx doesn't have direct page concept. This method converts to PDF first.
        
        Args:
            pages_per_split: Number of pages per split
            output_dir: Directory to save split files
            base_name: Base name for output files
            
        Returns:
            List of paths to created split files (PDFs)
        """
        # Import here to avoid circular dependency
        from services.format_converter import FormatConverter
        from services.pdf_operations import PDFSplitter
        
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Convert to PDF first
        converter = FormatConverter()
        temp_pdf = os.path.join(output_dir, "temp_conversion.pdf")
        
        try:
            print(f"Converting Word to PDF for splitting: {self.input_docx_path}")
            converter.word_to_pdf(self.input_docx_path, temp_pdf)
            
            # Now split the PDF
            splitter = PDFSplitter(temp_pdf)
            output_files = splitter.split_by_count(pages_per_split, output_dir, base_name)
            
            # Clean up temp PDF
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            
            return output_files
        
        except Exception as e:
            print(f"Error converting/splitting Word document: {e}")
            raise
    
    def split_by_names(self, names_file_path: str, pages_per_split: int, output_dir: str) -> List[str]:
        """
        Split Word document and name files according to a name list.
        Converts to PDF first for reliable splitting.
        
        Args:
            names_file_path: Path to Excel or TXT file with names
            pages_per_split: Number of pages per split file
            output_dir: Directory to save split files
            
        Returns:
            List of paths to created split files (PDFs)
        """
        # Import here to avoid circular dependency
        from services.format_converter import FormatConverter
        from services.pdf_operations import PDFSplitter
        
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Convert to PDF first
        converter = FormatConverter()
        temp_pdf = os.path.join(output_dir, "temp_conversion.pdf")
        
        try:
            print(f"Converting Word to PDF for named splitting: {self.input_docx_path}")
            converter.word_to_pdf(self.input_docx_path, temp_pdf)
            
            # Now split the PDF with names
            splitter = PDFSplitter(temp_pdf)
            output_files = splitter.split_by_names(names_file_path, pages_per_split, output_dir)
            
            # Clean up temp PDF
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            
            return output_files
        
        except Exception as e:
            print(f"Error converting/splitting Word document: {e}")
            raise


class WordMerger:
    """Handles Word document merging operations."""
    
    def merge_sequential(self, file_paths: List[str], output_path: str) -> str:
        """
        Merge multiple Word documents sequentially.
        
        Args:
            file_paths: List of Word file paths to merge
            output_path: Path for output merged document
            
        Returns:
            Path to created merged file
        """
        if not file_paths:
            raise ValueError("No files provided for merging")
        
        # Create new document from first file
        merged_doc = Document(file_paths[0])
        print(f"Starting merge with: {file_paths[0]}")
        
        # Append remaining documents
        for file_path in file_paths[1:]:
            print(f"Appending: {file_path}")
            self._append_document(merged_doc, file_path)
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Save merged document
        merged_doc.save(output_path)
        print(f"Created merged Word document: {output_path}")
        
        return output_path
    
    def merge_directory(self, directory_path: str, output_path: str) -> str:
        """
        Merge all Word documents in a directory sequentially (alphabetically sorted).
        
        Args:
            directory_path: Path to directory containing Word files
            output_path: Path for output merged document
            
        Returns:
            Path to created merged file
        """
        import glob
        
        # Get all Word files
        file_paths = []
        for ext in ['.docx', '.doc']:
            pattern = os.path.join(directory_path, f"*{ext}")
            file_paths.extend(glob.glob(pattern))
        
        file_paths = sorted(file_paths)
        
        if not file_paths:
            raise ValueError(f"No Word files found in {directory_path}")
        
        print(f"Found {len(file_paths)} files to merge in {directory_path}")
        return self.merge_sequential(file_paths, output_path)
    
    def _append_document(self, base_doc: Document, append_doc_path: str):
        """
        Append a Word document to the base document.
        
        Args:
            base_doc: Base document to append to
            append_doc_path: Path to document to append
        """
        append_doc = Document(append_doc_path)
        
        # Add page break before appending
        base_doc.add_page_break()
        
        # Copy all elements from append document
        for element in append_doc.element.body:
            if isinstance(element, CT_P):
                # Copy paragraph
                base_doc.element.body.append(element)
            elif isinstance(element, CT_Tbl):
                # Copy table
                base_doc.element.body.append(element)
    
    def merge_paired(self, file1_path: str, file2_path: str, output_path: str) -> str:
        """
        Merge two Word documents by interleaving pages.
        Note: Word doesn't have clear page boundaries. This converts to PDF first.
        
        Args:
            file1_path: Path to first Word document
            file2_path: Path to second Word document
            output_path: Path for output merged file
            
        Returns:
            Path to created merged file (PDF)
        """
        # Import here to avoid circular dependency
        from services.format_converter import FormatConverter
        from services.pdf_operations import PDFMerger
        
        # Convert both to PDF
        converter = FormatConverter()
        temp_dir = Path(output_path).parent / "temp_merge"
        temp_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            temp_pdf1 = str(temp_dir / "temp1.pdf")
            temp_pdf2 = str(temp_dir / "temp2.pdf")
            
            print(f"Converting Word documents to PDF for paired merge")
            converter.word_to_pdf(file1_path, temp_pdf1)
            converter.word_to_pdf(file2_path, temp_pdf2)
            
            # Merge PDFs in paired mode
            merger = PDFMerger()
            
            # Change output extension to PDF
            if output_path.lower().endswith('.docx'):
                output_path = output_path[:-5] + '.pdf'
            
            result = merger.merge_paired(temp_pdf1, temp_pdf2, output_path)
            
            # Clean up temp files
            import shutil
            shutil.rmtree(temp_dir)
            
            return result
        
        except Exception as e:
            print(f"Error in paired merge: {e}")
            raise
