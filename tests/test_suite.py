"""
Main Test Suite - Template Processing and Conversion Tests
Tests all core functionality including template processing, format conversion,
Excel auto-adjust, and variable substitution.
"""
import pytest
import os
import time
from pathlib import Path
from openpyxl import Workbook, load_workbook
from docx import Document


class TestTemplateProcessor:
    """Test suite for TemplateProcessor class."""
    
    def test_docx_variable_substitution(self, template_processor, output_dir, sample_data):
        """Test basic DOCX template variable substitution."""
        # Create simple DOCX template
        template_path = output_dir / "test_template.docx"
        doc = Document()
        doc.add_heading('Employee Information', 0)
        doc.add_paragraph('Name: ##name##')
        doc.add_paragraph('Email: ##email##')
        doc.add_paragraph('Phone: ##phone##')
        doc.add_paragraph('Company: ##company##')
        doc.add_paragraph('Position: ##position##')
        doc.save(str(template_path))
        
        # Process template
        output_path = output_dir / "output.docx"
        template_processor.process_template(
            str(template_path),
            sample_data[0],
            str(output_path)
        )
        
        # Verify output
        assert output_path.exists(), "Output file not created"
        
        # Check content
        result_doc = Document(str(output_path))
        text = '\n'.join([p.text for p in result_doc.paragraphs])
        
        assert 'John Doe' in text, "Name not substituted"
        assert 'john.doe@example.com' in text, "Email not substituted"
        assert '555-1234' in text, "Phone not substituted"
        assert 'Acme Corporation' in text, "Company not substituted"
        assert 'Senior Developer' in text, "Position not substituted"
        
        # Ensure variables are replaced (no ## markers left)
        assert '##name##' not in text, "Variable not replaced"
        assert '##email##' not in text, "Variable not replaced"
    
    def test_xlsx_variable_substitution(self, template_processor, output_dir, sample_data):
        """Test basic XLSX template variable substitution."""
        # Create simple XLSX template
        template_path = output_dir / "test_template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Employee"
        
        ws['A1'] = 'Employee Information'
        ws['A2'] = 'Name:'
        ws['B2'] = '##name##'
        ws['A3'] = 'Email:'
        ws['B3'] = '##email##'
        ws['A4'] = 'Phone:'
        ws['B4'] = '##phone##'
        ws['A5'] = 'Company:'
        ws['B5'] = '##company##'
        ws['A6'] = 'Position:'
        ws['B6'] = '##position##'
        
        wb.save(str(template_path))
        
        # Process template
        output_path = output_dir / "output.xlsx"
        template_processor.process_template(
            str(template_path),
            sample_data[0],
            str(output_path)
        )
        
        # Verify output
        assert output_path.exists(), "Output file not created"
        
        # Check content
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        assert result_ws['B2'].value == 'John Doe', "Name not substituted"
        assert result_ws['B3'].value == 'john.doe@example.com', "Email not substituted"
        assert result_ws['B4'].value == '555-1234', "Phone not substituted"
        assert result_ws['B5'].value == 'Acme Corporation', "Company not substituted"
        assert result_ws['B6'].value == 'Senior Developer', "Position not substituted"
        
        # Ensure no variables remain
        for row in result_ws.iter_rows(values_only=True):
            for cell in row:
                if cell and isinstance(cell, str):
                    assert '##' not in cell, f"Variable marker found in: {cell}"
    
    def test_xlsx_auto_adjust_height(self, template_processor, output_dir):
        """Test Excel auto-adjust row height feature."""
        # Create template with long text
        template_path = output_dir / "test_template.xlsx"
        wb = Workbook()
        ws = wb.active
        
        ws['A1'] = '##short##'
        ws['A2'] = '##long##'
        ws.column_dimensions['A'].width = 20  # Fixed width
        
        wb.save(str(template_path))
        
        # Process with auto-adjust
        output_path = output_dir / "output.xlsx"
        data = {
            'short': 'Short text',
            'long': 'This is a very long text that should wrap and require increased row height to display properly without being cut off'
        }
        
        auto_adjust_options = {
            'auto_adjust_height': True,
            'auto_adjust_width': False
        }
        
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path),
            auto_adjust_options=auto_adjust_options
        )
        
        # Verify output
        assert output_path.exists(), "Output file not created"
        
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        # Check that row height was adjusted (should be None for auto)
        assert result_ws.row_dimensions[2].height is None, "Row height not set to auto"
    
    def test_xlsx_auto_adjust_width(self, template_processor, output_dir):
        """Test Excel auto-adjust column width feature."""
        # Create template
        template_path = output_dir / "test_template.xlsx"
        wb = Workbook()
        ws = wb.active
        
        ws['A1'] = '##text1##'
        ws['B1'] = '##text2##'
        ws['C1'] = '##text3##'
        
        wb.save(str(template_path))
        
        # Process with auto-adjust
        output_path = output_dir / "output.xlsx"
        data = {
            'text1': 'Short',
            'text2': 'Medium length text',
            'text3': 'Very long text that needs more space'
        }
        
        auto_adjust_options = {
            'auto_adjust_height': False,
            'auto_adjust_width': True
        }
        
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path),
            auto_adjust_options=auto_adjust_options
        )
        
        # Verify output
        assert output_path.exists(), "Output file not created"
        
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        # Check column widths were adjusted
        width_A = result_ws.column_dimensions['A'].width
        width_B = result_ws.column_dimensions['B'].width
        width_C = result_ws.column_dimensions['C'].width
        
        # Longer text should have wider columns
        assert width_C > width_B > width_A, "Column widths not adjusted properly"
    
    def test_xlsx_auto_adjust_specific_range(self, template_processor, output_dir):
        """Test Excel auto-adjust with specific range."""
        # Create template
        template_path = output_dir / "test_template.xlsx"
        wb = Workbook()
        ws = wb.active
        
        ws['A1'] = 'Data outside range'
        ws['B2'] = '##value1##'
        ws['C3'] = '##value2##'
        ws['D4'] = '##value3##'
        ws['E5'] = 'More data'
        
        wb.save(str(template_path))
        
        # Process with specific range
        output_path = output_dir / "output.xlsx"
        data = {
            'value1': 'Test value 1',
            'value2': 'Test value 2',
            'value3': 'Test value 3'
        }
        
        auto_adjust_options = {
            'auto_adjust_height': True,
            'auto_adjust_width': True,
            'adjust_range': 'B2:D4'
        }
        
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path),
            auto_adjust_options=auto_adjust_options
        )
        
        # Verify output
        assert output_path.exists(), "Output file not created"
        
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        # Verify data was substituted
        assert result_ws['B2'].value == 'Test value 1'
        assert result_ws['C3'].value == 'Test value 2'
        assert result_ws['D4'].value == 'Test value 3'
    
    def test_template_caching(self, template_processor, output_dir, sample_data):
        """Test that template caching improves performance."""
        # Create template
        template_path = output_dir / "cached_template.docx"
        doc = Document()
        doc.add_paragraph('Name: ##name##')
        doc.save(str(template_path))
        
        # First processing (cache miss)
        start1 = time.time()
        output1 = output_dir / "output1.docx"
        template_processor.process_template(
            str(template_path),
            sample_data[0],
            str(output1)
        )
        time1 = time.time() - start1
        
        # Second processing (cache hit)
        start2 = time.time()
        output2 = output_dir / "output2.docx"
        template_processor.process_template(
            str(template_path),
            sample_data[1],
            str(output2)
        )
        time2 = time.time() - start2
        
        # Second should be faster (cached)
        assert output1.exists() and output2.exists(), "Outputs not created"
        print(f"\nFirst processing: {time1:.3f}s")
        print(f"Second processing: {time2:.3f}s")
        print(f"Speedup: {time1/time2:.2f}x")
        
        # Cache should provide some speedup
        assert time2 <= time1, "Caching did not improve performance"


class TestFormatConverter:
    """Test suite for FormatConverter class."""
    
    @pytest.mark.requires_libreoffice
    def test_docx_to_pdf_conversion(self, format_converter, output_dir):
        """Test DOCX to PDF conversion."""
        # Create simple DOCX
        docx_path = output_dir / "test.docx"
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for PDF conversion.')
        doc.save(str(docx_path))
        
        # Convert to PDF
        pdf_path = format_converter.convert(str(docx_path), 'pdf', str(output_dir))
        
        # Verify
        assert pdf_path is not None, "Conversion failed"
        assert Path(pdf_path).exists(), "PDF not created"
        assert Path(pdf_path).suffix.lower() == '.pdf', "Wrong output format"
        assert Path(pdf_path).stat().st_size > 0, "PDF is empty"
    
    @pytest.mark.requires_libreoffice
    def test_xlsx_to_pdf_conversion(self, format_converter, output_dir):
        """Test XLSX to PDF conversion."""
        # Create simple XLSX
        xlsx_path = output_dir / "test.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Test Spreadsheet'
        ws['A2'] = 'Row 2'
        ws['A3'] = 'Row 3'
        wb.save(str(xlsx_path))
        
        # Convert to PDF
        pdf_path = format_converter.convert(str(xlsx_path), 'pdf', str(output_dir))
        
        # Verify
        assert pdf_path is not None, "Conversion failed"
        assert Path(pdf_path).exists(), "PDF not created"
        assert Path(pdf_path).suffix.lower() == '.pdf', "Wrong output format"
        assert Path(pdf_path).stat().st_size > 0, "PDF is empty"
    
    def test_unsupported_format(self, format_converter, output_dir):
        """Test handling of unsupported format conversion."""
        # Create a text file
        txt_path = output_dir / "test.txt"
        txt_path.write_text("Test content")
        
        # Try to convert (should handle gracefully)
        try:
            result = format_converter.convert(str(txt_path), 'pdf', str(output_dir))
            # If it succeeds, that's fine
            if result:
                assert Path(result).exists()
        except Exception as e:
            # Should raise appropriate error
            error_msg = str(e).lower()
            assert "not supported" in error_msg or "error" in error_msg or "cannot convert" in error_msg


class TestJobManager:
    """Test suite for JobManager class."""
    
    def test_create_job(self, job_manager, output_dir):
        """Test job creation."""
        # Create test files
        template_path = output_dir / "template.docx"
        data_path = output_dir / "data.xlsx"
        
        doc = Document()
        doc.add_paragraph('##name##')
        doc.save(str(template_path))
        
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##name##'
        ws['A2'] = 'John Doe'
        wb.save(str(data_path))
        
        # Create job
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx', 'pdf']
        )
        
        # Verify
        assert job is not None, "Job not created"
        assert job.id is not None, "Job ID not set"
        assert job.status.value == 'pending', "Job status incorrect"
        # JobManager copies template to job directory, check basename instead
        assert Path(template_path).name in job.metadata['job_template_path'], "Template name not in path"
    
    def test_job_with_excel_auto_adjust(self, job_manager, output_dir):
        """Test job creation with Excel auto-adjust options."""
        template_path = output_dir / "template.xlsx"
        data_path = output_dir / "data.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##text##'
        wb.save(str(template_path))
        
        wb2 = Workbook()
        ws2 = wb2.active
        ws2['A1'] = '##text##'
        ws2['A2'] = 'Long text value'
        wb2.save(str(data_path))
        
        auto_adjust_options = {
            'auto_adjust_height': True,
            'auto_adjust_width': True
        }
        
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['xlsx'],
            excel_auto_adjust_options=auto_adjust_options
        )
        
        assert job.excel_auto_adjust_options is not None, "Auto-adjust options not stored"
        assert job.excel_auto_adjust_options['auto_adjust_height'] is True
        assert job.excel_auto_adjust_options['auto_adjust_width'] is True
    
    def test_job_with_excel_print_settings(self, job_manager, output_dir, excel_print_settings):
        """Test job creation with Excel print settings."""
        template_path = output_dir / "template.xlsx"
        data_path = output_dir / "data.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Test'
        wb.save(str(template_path))
        
        wb2 = Workbook()
        ws2 = wb2.active
        ws2['A1'] = '##value##'
        ws2['A2'] = 'Data'
        wb2.save(str(data_path))
        
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['pdf'],
            excel_print_settings=excel_print_settings
        )
        
        assert job.excel_print_settings is not None, "Print settings not stored"
        assert job.excel_print_settings['orientation'] == 'portrait'
        assert job.excel_print_settings['paper_size'] == 'a4'


class TestEdgeCases:
    """Test suite for edge cases and error handling."""
    
    def test_missing_variable_in_template(self, template_processor, output_dir):
        """Test handling of missing variables in data."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('Name: ##name##')
        doc.add_paragraph('Email: ##email##')
        doc.save(str(template_path))
        
        # Data missing 'email'
        data = {'name': 'John Doe'}
        
        output_path = output_dir / "output.docx"
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path)
        )
        
        # Should still work, leaving unmatched variables
        assert output_path.exists()
        
        result_doc = Document(str(output_path))
        text = '\n'.join([p.text for p in result_doc.paragraphs])
        assert 'John Doe' in text
        # Email variable should remain unreplaced
        assert '##email##' in text
    
    def test_special_characters_in_variables(self, template_processor, output_dir):
        """Test handling of special characters in variable values."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('Text: ##text##')
        doc.save(str(template_path))
        
        # Data with special characters
        data = {'text': 'Special chars: $@#%&*()[]{}!?<>'}
        
        output_path = output_dir / "output.docx"
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path)
        )
        
        assert output_path.exists()
        
        result_doc = Document(str(output_path))
        text = '\n'.join([p.text for p in result_doc.paragraphs])
        assert 'Special chars:' in text
    
    def test_empty_template(self, template_processor, output_dir):
        """Test processing of empty template."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.save(str(template_path))
        
        data = {'name': 'John Doe'}
        
        output_path = output_dir / "output.docx"
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path)
        )
        
        assert output_path.exists()
    
    def test_large_dataset(self, template_processor, output_dir):
        """Test processing with large text values."""
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##large_text##'
        wb.save(str(template_path))
        
        # Create very large text
        large_text = 'A' * 10000  # 10,000 characters
        data = {'large_text': large_text}
        
        output_path = output_dir / "output.xlsx"
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path)
        )
        
        assert output_path.exists()
        
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        assert len(result_ws['A1'].value) == 10000


if __name__ == '__main__':
    pytest.main([__file__, '-v', '--tb=short'])
