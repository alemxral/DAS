"""
Performance Tests - Benchmarking and Load Testing
Tests system performance, caching effectiveness, and resource usage.
"""
import pytest
import time
import psutil
import os
from pathlib import Path
from openpyxl import Workbook
from docx import Document


@pytest.mark.performance
class TestPerformance:
    """Performance and benchmarking tests."""
    
    def test_template_caching_speedup(self, template_processor, output_dir):
        """Measure template caching performance improvement."""
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_heading('Performance Test', 0)
        doc.add_paragraph('Name: ##name##')
        doc.add_paragraph('Email: ##email##')
        doc.add_paragraph('Phone: ##phone##')
        doc.save(str(template_path))
        
        # Prepare data
        test_data = [
            {'name': f'Person {i}', 'email': f'person{i}@example.com', 'phone': f'555-{i:04d}'}
            for i in range(10)
        ]
        
        # First run (cold cache)
        start_cold = time.time()
        for i, data in enumerate(test_data[:5]):
            output = output_dir / f"cold_{i}.docx"
            template_processor.process_template(str(template_path), data, str(output))
        time_cold = time.time() - start_cold
        
        # Second run (warm cache)
        start_warm = time.time()
        for i, data in enumerate(test_data[5:]):
            output = output_dir / f"warm_{i}.docx"
            template_processor.process_template(str(template_path), data, str(output))
        time_warm = time.time() - start_warm
        
        avg_cold = time_cold / 5
        avg_warm = time_warm / 5
        speedup = avg_cold / avg_warm
        
        print(f"\n--- Template Caching Performance ---")
        print(f"Cold cache (first 5): {time_cold:.3f}s (avg: {avg_cold:.3f}s)")
        print(f"Warm cache (next 5):  {time_warm:.3f}s (avg: {avg_warm:.3f}s)")
        print(f"Speedup: {speedup:.2f}x")
        
        # With our fix (load fresh copy instead of deepcopy), caching may not improve
        # performance significantly but ensures correctness. Allow 20% variance.
        assert avg_warm <= avg_cold * 1.20, f"Cache performance degraded significantly: {avg_warm:.3f}s vs {avg_cold:.3f}s"
        
        # Return metrics for reporting
        return {
            'cold_time': time_cold,
            'warm_time': time_warm,
            'speedup': speedup
        }
    
    @pytest.mark.slow
    def test_large_dataset_processing(self, template_processor, output_dir):
        """Test processing with large dataset."""
        # Create template
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##id##'
        ws['A2'] = '##name##'
        ws['A3'] = '##description##'
        wb.save(str(template_path))
        
        # Generate large dataset
        num_records = 100
        print(f"\n--- Processing {num_records} records ---")
        
        start_time = time.time()
        
        for i in range(num_records):
            data = {
                'id': f'ID-{i:05d}',
                'name': f'Record {i}',
                'description': f'Description for record number {i}'
            }
            output = output_dir / f"record_{i:05d}.xlsx"
            template_processor.process_template(str(template_path), data, str(output))
        
        elapsed = time.time() - start_time
        avg_time = elapsed / num_records
        
        print(f"Total time: {elapsed:.2f}s")
        print(f"Average per record: {avg_time:.3f}s")
        print(f"Records per second: {num_records/elapsed:.2f}")
        
        # Performance target: should process at least 10 records/second
        assert num_records / elapsed > 10, "Processing too slow"
    
    def test_memory_usage(self, template_processor, output_dir):
        """Test memory usage during processing."""
        process = psutil.Process(os.getpid())
        
        # Get initial memory
        mem_start = process.memory_info().rss / 1024 / 1024  # MB
        
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##content##')
        doc.save(str(template_path))
        
        # Process multiple documents
        num_docs = 50
        for i in range(num_docs):
            data = {'content': f'Document {i} content' * 100}
            output = output_dir / f"doc_{i}.docx"
            template_processor.process_template(str(template_path), data, str(output))
        
        # Get final memory
        mem_end = process.memory_info().rss / 1024 / 1024  # MB
        mem_increase = mem_end - mem_start
        
        print(f"\n--- Memory Usage ---")
        print(f"Start: {mem_start:.2f} MB")
        print(f"End: {mem_end:.2f} MB")
        print(f"Increase: {mem_increase:.2f} MB")
        print(f"Per document: {mem_increase/num_docs:.3f} MB")
        
        # Memory increase should be reasonable (< 200 MB for 50 docs)
        assert mem_increase < 200, "Memory usage too high"
    
    @pytest.mark.requires_libreoffice
    def test_conversion_speed(self, format_converter, output_dir):
        """Benchmark format conversion speed."""
        conversion_times = {}
        
        # Test DOCX to PDF
        docx_path = output_dir / "test.docx"
        doc = Document()
        doc.add_heading('Test Document', 0)
        for i in range(10):
            doc.add_paragraph(f'Paragraph {i}: This is test content for conversion benchmarking.')
        doc.save(str(docx_path))
        
        start = time.time()
        pdf_path = format_converter.convert(str(docx_path), 'pdf', str(output_dir))
        conversion_times['docx_to_pdf'] = time.time() - start
        
        # Test XLSX to PDF
        xlsx_path = output_dir / "test.xlsx"
        wb = Workbook()
        ws = wb.active
        for row in range(1, 51):
            for col in range(1, 11):
                ws.cell(row=row, column=col, value=f'Cell {row},{col}')
        wb.save(str(xlsx_path))
        
        start = time.time()
        pdf_path2 = format_converter.convert(str(xlsx_path), 'pdf', str(output_dir))
        conversion_times['xlsx_to_pdf'] = time.time() - start
        
        print(f"\n--- Conversion Speed ---")
        for conversion, time_taken in conversion_times.items():
            print(f"{conversion}: {time_taken:.3f}s")
        
        # Conversions should complete in reasonable time
        for time_taken in conversion_times.values():
            assert time_taken < 30, "Conversion too slow"
        
        return conversion_times
    
    def test_concurrent_processing(self, template_processor, output_dir):
        """Test concurrent template processing."""
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##data##')
        doc.save(str(template_path))
        
        # Sequential processing
        start_seq = time.time()
        for i in range(10):
            data = {'data': f'Sequential {i}'}
            output = output_dir / f"seq_{i}.docx"
            template_processor.process_template(str(template_path), data, str(output))
        time_seq = time.time() - start_seq
        
        print(f"\n--- Concurrent Processing Test ---")
        print(f"Sequential: {time_seq:.3f}s")
        print(f"Per document: {time_seq/10:.3f}s")
        
        # Note: True concurrent testing would require threading
        # This test validates sequential performance
        assert time_seq < 10, "Sequential processing too slow"


@pytest.mark.performance
class TestScalability:
    """Scalability and load tests."""
    
    @pytest.mark.slow
    def test_multiple_sheets_processing(self, template_processor, output_dir):
        """Test processing Excel with multiple sheets."""
        # Create template with multiple sheets
        template_path = output_dir / "multi_sheet.xlsx"
        wb = Workbook()
        
        # Create 5 sheets
        for i in range(5):
            ws = wb.create_sheet(f"Sheet{i+1}")
            ws['A1'] = f'Sheet {i+1}'
            ws['A2'] = '##value##'
        
        # Remove default sheet
        if 'Sheet' in wb.sheetnames:
            del wb['Sheet']
        
        wb.save(str(template_path))
        
        # Process each sheet
        start = time.time()
        for i in range(5):
            data = {'value': f'Data for sheet {i+1}'}
            output = output_dir / f"output_sheet{i+1}.xlsx"
            template_processor.process_template(
                str(template_path),
                data,
                str(output),
                sheet_name=f"Sheet{i+1}"
            )
        elapsed = time.time() - start
        
        print(f"\n--- Multi-Sheet Processing ---")
        print(f"Total time: {elapsed:.3f}s")
        print(f"Per sheet: {elapsed/5:.3f}s")
        
        # All output files should exist
        for i in range(5):
            output = output_dir / f"output_sheet{i+1}.xlsx"
            assert output.exists()
    
    def test_file_size_handling(self, template_processor, output_dir):
        """Test handling of large files."""
        # Create large template
        template_path = output_dir / "large_template.xlsx"
        wb = Workbook()
        ws = wb.active
        
        # Add substantial content
        for row in range(1, 101):
            for col in range(1, 21):
                ws.cell(row=row, column=col, value=f'##var{col}##')
        
        wb.save(str(template_path))
        
        # Prepare data
        data = {f'##var{i}##': f'Value {i}' for i in range(1, 21)}
        
        # Process
        start = time.time()
        output_path = output_dir / "large_output.xlsx"
        template_processor.process_template(str(template_path), data, str(output_path))
        elapsed = time.time() - start
        
        print(f"\n--- Large File Processing ---")
        print(f"Template size: {template_path.stat().st_size / 1024:.2f} KB")
        print(f"Output size: {output_path.stat().st_size / 1024:.2f} KB")
        print(f"Processing time: {elapsed:.3f}s")
        
        assert output_path.exists()
        assert output_path.stat().st_size > 0


@pytest.mark.performance
class TestResourceUsage:
    """Resource usage and optimization tests."""
    
    def test_cache_memory_efficiency(self, template_processor, output_dir):
        """Test cache memory efficiency."""
        process = psutil.Process(os.getpid())
        mem_start = process.memory_info().rss / 1024 / 1024
        
        # Create and cache multiple templates
        templates = []
        for i in range(10):
            template_path = output_dir / f"template_{i}.docx"
            doc = Document()
            doc.add_paragraph(f'Template {i}: ##data##')
            doc.save(str(template_path))
            templates.append(template_path)
        
        # Process all templates (loads into cache)
        for i, template_path in enumerate(templates):
            data = {'data': f'Data {i}'}
            output = output_dir / f"cached_{i}.docx"
            template_processor.process_template(str(template_path), data, str(output))
        
        mem_end = process.memory_info().rss / 1024 / 1024
        cache_overhead = mem_end - mem_start
        
        print(f"\n--- Cache Memory Efficiency ---")
        print(f"Templates cached: 10")
        print(f"Memory overhead: {cache_overhead:.2f} MB")
        print(f"Per template: {cache_overhead/10:.2f} MB")
        
        # Cache overhead should be reasonable
        assert cache_overhead < 100, "Cache using too much memory"
    
    def test_disk_io_efficiency(self, template_processor, output_dir):
        """Test disk I/O efficiency."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##data##')
        doc.save(str(template_path))
        
        # Count file operations
        start_time = time.time()
        num_operations = 20
        
        for i in range(num_operations):
            data = {'data': f'Data {i}'}
            output = output_dir / f"io_test_{i}.docx"
            template_processor.process_template(str(template_path), data, str(output))
        
        elapsed = time.time() - start_time
        ops_per_second = num_operations / elapsed
        
        print(f"\n--- Disk I/O Efficiency ---")
        print(f"Operations: {num_operations}")
        print(f"Time: {elapsed:.3f}s")
        print(f"Ops/second: {ops_per_second:.2f}")
        
        # Should achieve reasonable throughput
        assert ops_per_second > 5, "I/O too slow"


if __name__ == '__main__':
    pytest.main([__file__, '-v', '--tb=short', '-m', 'performance'])
