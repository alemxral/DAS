"""
Integration Tests - End-to-End Workflow Testing
Tests complete workflows from job creation to output generation.
"""
import pytest
import os
import time
import zipfile
from pathlib import Path
from openpyxl import Workbook, load_workbook
from docx import Document


@pytest.mark.integration
class TestEndToEndWorkflows:
    """Integration tests for complete workflows."""
    
    def test_single_template_single_output(self, job_manager, output_dir):
        """Test complete workflow: single template, single data row."""
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_heading('Employee Report', 0)
        doc.add_paragraph('Name: ##name##')
        doc.add_paragraph('Email: ##email##')
        doc.save(str(template_path))
        
        # Create data
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['name', 'email', 'filename'])
        ws.append(['John Doe', 'john@example.com', 'john_report'])
        wb.save(str(data_path))
        
        # Create and process job
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx']
        )
        
        assert job is not None
        assert job.status.value == 'pending'
        
        # Process job
        processed_job = job_manager.process_job(job.id)
        
        # Verify results
        assert processed_job.status.value == 'completed'
        assert len(processed_job.output_files) > 0
        
        # Check output file exists
        output_file = processed_job.output_files[0]
        assert Path(output_file).exists()
    
    def test_single_template_multiple_outputs(self, job_manager, output_dir):
        """Test complete workflow: single template, multiple data rows."""
        # Create template
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Name:'
        ws['B1'] = '##name##'
        ws['A2'] = 'Department:'
        ws['B2'] = '##department##'
        wb.save(str(template_path))
        
        # Create data with multiple rows
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['name', 'department', 'filename'])
        ws.append(['John Doe', 'Engineering', 'john'])
        ws.append(['Jane Smith', 'Marketing', 'jane'])
        ws.append(['Bob Johnson', 'Sales', 'bob'])
        wb.save(str(data_path))
        
        # Create and process job
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['xlsx']
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify results
        assert processed_job.status.value == 'completed'
        assert len(processed_job.output_files) == 3  # 3 data rows
        
        # Verify all files exist
        for output_file in processed_job.output_files:
            assert Path(output_file).exists()
    
    @pytest.mark.requires_libreoffice
    def test_multiple_output_formats(self, job_manager, output_dir):
        """Test complete workflow with multiple output formats."""
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('Name: ##name##')
        doc.save(str(template_path))
        
        # Create data
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['##name##', '##filename##'])
        ws.append(['John Doe', 'output'])
        wb.save(str(data_path))
        
        # Create job with multiple formats
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx', 'pdf']
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify results
        assert processed_job.status.value == 'completed'
        
        # Should have both DOCX and PDF
        output_files = processed_job.output_files
        docx_files = [f for f in output_files if f.endswith('.docx')]
        pdf_files = [f for f in output_files if f.endswith('.pdf')]
        
        assert len(docx_files) > 0, "No DOCX files generated"
        assert len(pdf_files) > 0, "No PDF files generated"
    
    def test_excel_auto_adjust_integration(self, job_manager, output_dir):
        """Test Excel auto-adjust in complete workflow."""
        # Create template with cells that need adjustment
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##short##'
        ws['A2'] = '##long##'
        ws.column_dimensions['A'].width = 15
        wb.save(str(template_path))
        
        # Create data
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['short', 'long', 'filename'])
        ws.append(['Hi', 'This is a very long text that needs width adjustment', 'output'])
        wb.save(str(data_path))
        
        # Create job with auto-adjust
        auto_adjust_options = {
            'auto_adjust_height': True,
            'auto_adjust_width': True
        }
        
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['xlsx'],
            excel_auto_adjust_options=auto_adjust_options
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify
        assert processed_job.status.value == 'completed'
        assert len(processed_job.output_files) > 0
        
        # Check that output has adjusted dimensions
        output_file = processed_job.output_files[0]
        result_wb = load_workbook(output_file)
        result_ws = result_wb.active
        
        # Column A should be wider than original 15
        assert result_ws.column_dimensions['A'].width > 15
    
    def test_zip_archive_creation(self, job_manager, output_dir):
        """Test that ZIP archive is created for job outputs."""
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##name##')
        doc.save(str(template_path))
        
        # Create data with multiple rows
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['##name##', '##filename##'])
        ws.append(['Person 1', 'file1'])
        ws.append(['Person 2', 'file2'])
        wb.save(str(data_path))
        
        # Process job
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx']
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify ZIP created
        assert processed_job.zip_file_path is not None
        assert Path(processed_job.zip_file_path).exists()
        assert Path(processed_job.zip_file_path).suffix == '.zip'
        
        # Verify ZIP contains output files
        with zipfile.ZipFile(processed_job.zip_file_path, 'r') as zip_file:
            file_list = zip_file.namelist()
            assert len(file_list) > 0
            
            # Should contain the output files
            docx_files = [f for f in file_list if f.endswith('.docx')]
            assert len(docx_files) == 2  # 2 data rows
    
    def test_custom_output_directory(self, job_manager, output_dir, tmp_path):
        """Test job with custom output directory."""
        custom_output = tmp_path / "custom_output"
        custom_output.mkdir()
        
        # Create template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##name##')
        doc.save(str(template_path))
        
        # Create data
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['##name##', '##filename##'])
        ws.append(['John Doe', 'output'])
        wb.save(str(data_path))
        
        # Create job with custom output directory
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx'],
            output_directory=str(custom_output)
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify output in custom directory
        assert processed_job.status.value == 'completed'
        
        # Check ZIP file is in custom directory (JobManager copies ZIP, not individual files)
        zip_files = list(custom_output.glob('*.zip'))
        assert len(zip_files) > 0, "ZIP file not found in custom output directory"
        
        # Verify ZIP contains output files
        import zipfile
        with zipfile.ZipFile(zip_files[0], 'r') as zip_file:
            file_list = zip_file.namelist()
            output_files = [f for f in file_list if f.endswith('.docx')]
            assert len(output_files) > 0


@pytest.mark.integration
class TestMultiTemplateWorkflows:
    """Integration tests for multi-template mode."""
    
    def test_multi_template_basic(self, job_manager, output_dir):
        """Test basic multi-template workflow."""
        # Create multiple templates
        template1_path = output_dir / "template1.docx"
        doc1 = Document()
        doc1.add_heading('Template 1', 0)
        doc1.add_paragraph('Name: ##name##')
        doc1.save(str(template1_path))
        
        template2_path = output_dir / "template2.docx"
        doc2 = Document()
        doc2.add_heading('Template 2', 0)
        doc2.add_paragraph('Email: ##email##')
        doc2.save(str(template2_path))
        
        # Create data
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['name', 'email', 'filename'])
        ws.append(['John Doe', 'john@example.com', 'john'])
        wb.save(str(data_path))
        
        # Create job with multiple templates
        templates = [
            {'path': str(template1_path), 'priority': 1},
            {'path': str(template2_path), 'priority': 2}
        ]
        
        job = job_manager.create_job(
            template_path=None,
            data_path=str(data_path),
            output_formats=['docx'],
            templates=templates
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify results
        assert processed_job.status.value == 'completed'
        # Should have 2 output files (1 data row × 2 templates)
        assert len(processed_job.output_files) >= 2
    
    def test_multi_template_with_excel_sheets(self, job_manager, output_dir):
        """Test multi-template with different Excel sheets."""
        # Create templates
        template1_path = output_dir / "template1.xlsx"
        wb1 = Workbook()
        ws1 = wb1.active
        ws1.title = "Report"
        ws1['A1'] = '##value1##'
        wb1.save(str(template1_path))
        
        template2_path = output_dir / "template2.xlsx"
        wb2 = Workbook()
        ws2 = wb2.active
        ws2.title = "Summary"
        ws2['A1'] = '##value2##'
        wb2.save(str(template2_path))
        
        # Create data
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['value1', 'value2', 'filename'])
        ws.append(['Data 1', 'Data 2', 'output'])
        wb.save(str(data_path))
        
        # Create job
        templates = [
            {'path': str(template1_path), 'priority': 1, 'sheet': 'Report'},
            {'path': str(template2_path), 'priority': 2, 'sheet': 'Summary'}
        ]
        
        job = job_manager.create_job(
            template_path=None,
            data_path=str(data_path),
            output_formats=['xlsx'],
            templates=templates
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify
        assert processed_job.status.value == 'completed'
        assert len(processed_job.output_files) >= 2


@pytest.mark.integration
class TestErrorRecovery:
    """Integration tests for error handling and recovery."""
    
    def test_invalid_template_path(self, job_manager, output_dir):
        """Test handling of invalid template path."""
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['##name##'])
        ws.append(['John'])
        wb.save(str(data_path))
        
        # Create job with non-existent template
        job = job_manager.create_job(
            template_path="c:\\nonexistent\\template.docx",
            data_path=str(data_path),
            output_formats=['docx']
        )
        
        # Process should fail gracefully
        processed_job = job_manager.process_job(job.id)
        
        assert processed_job.status.value == 'failed'
        assert processed_job.error_message is not None
    
    def test_invalid_data_path(self, job_manager, output_dir):
        """Test handling of invalid data path."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##name##')
        doc.save(str(template_path))
        
        # Create job with non-existent data
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path="c:\\nonexistent\\data.xlsx",
            output_formats=['docx']
        )
        
        # Process should fail gracefully
        processed_job = job_manager.process_job(job.id)
        
        assert processed_job.status.value == 'failed'
        assert processed_job.error_message is not None
    
    def test_corrupted_template(self, job_manager, output_dir):
        """Test handling of corrupted template file."""
        # Create corrupted file
        template_path = output_dir / "corrupted.docx"
        template_path.write_text("This is not a valid DOCX file")
        
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['##name##'])
        ws.append(['John'])
        wb.save(str(data_path))
        
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx']
        )
        
        # Process should fail gracefully
        processed_job = job_manager.process_job(job.id)
        
        assert processed_job.status.value == 'failed'


if __name__ == '__main__':
    pytest.main([__file__, '-v', '--tb=short'])
