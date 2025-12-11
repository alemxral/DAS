"""
Validation Tests - Output File Integrity and Correctness
Tests output file validity, data integrity, and feature application.
"""
import pytest
from pathlib import Path
from openpyxl import load_workbook
from docx import Document
import zipfile
import re


class TestOutputValidation:
    """Tests for validating output file integrity."""
    
    def test_docx_file_validity(self, template_processor, output_dir):
        """Test that generated DOCX files are valid."""
        # Create and process template
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('Test: ##value##')
        doc.save(str(template_path))
        
        output_path = output_dir / "output.docx"
        template_processor.process_template(
            str(template_path),
            {'value': 'Valid Document'},
            str(output_path)
        )
        
        # Verify file is valid DOCX
        assert output_path.exists(), "Output file not created"
        assert output_path.stat().st_size > 0, "Output file is empty"
        
        # Try to open as DOCX
        try:
            result_doc = Document(str(output_path))
            paragraphs = [p.text for p in result_doc.paragraphs]
            assert len(paragraphs) > 0, "No paragraphs in document"
            assert any('Valid Document' in p for p in paragraphs), "Content not found"
        except Exception as e:
            pytest.fail(f"DOCX file is invalid: {e}")
    
    def test_xlsx_file_validity(self, template_processor, output_dir):
        """Test that generated XLSX files are valid."""
        # Create and process template
        from openpyxl import Workbook
        
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##value##'
        wb.save(str(template_path))
        
        output_path = output_dir / "output.xlsx"
        template_processor.process_template(
            str(template_path),
            {'value': 'Valid Spreadsheet'},
            str(output_path)
        )
        
        # Verify file is valid XLSX
        assert output_path.exists(), "Output file not created"
        assert output_path.stat().st_size > 0, "Output file is empty"
        
        # Try to open as XLSX
        try:
            result_wb = load_workbook(str(output_path))
            result_ws = result_wb.active
            assert result_ws['A1'].value == 'Valid Spreadsheet', "Content not correct"
        except Exception as e:
            pytest.fail(f"XLSX file is invalid: {e}")
    
    @pytest.mark.requires_libreoffice
    def test_pdf_file_validity(self, format_converter, output_dir):
        """Test that generated PDF files are valid."""
        # Create source document
        from docx import Document
        
        docx_path = output_dir / "source.docx"
        doc = Document()
        doc.add_paragraph('PDF Test Content')
        doc.save(str(docx_path))
        
        # Convert to PDF
        pdf_path = format_converter.convert(str(docx_path), 'pdf', str(output_dir))
        
        # Verify PDF
        assert pdf_path is not None, "PDF conversion failed"
        assert Path(pdf_path).exists(), "PDF file not created"
        assert Path(pdf_path).stat().st_size > 0, "PDF file is empty"
        
        # Check PDF header
        with open(pdf_path, 'rb') as f:
            header = f.read(4)
            assert header == b'%PDF', "Not a valid PDF file"


class TestVariableSubstitution:
    """Tests for verifying variable substitution correctness."""
    
    def test_all_variables_replaced_docx(self, template_processor, output_dir):
        """Test that all variables are replaced in DOCX."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##var1## and ##var2## and ##var3##')
        doc.save(str(template_path))
        
        output_path = output_dir / "output.docx"
        data = {
            'var1': 'Value1',
            'var2': 'Value2',
            'var3': 'Value3'
        }
        
        template_processor.process_template(str(template_path), data, str(output_path))
        
        # Check no ## markers remain
        result_doc = Document(str(output_path))
        text = '\n'.join([p.text for p in result_doc.paragraphs])
        
        # Find any remaining variable markers
        remaining_vars = re.findall(r'##[^#]+##', text)
        assert len(remaining_vars) == 0, f"Variables not replaced: {remaining_vars}"
        
        # Verify values present
        assert 'Value1' in text
        assert 'Value2' in text
        assert 'Value3' in text
    
    def test_all_variables_replaced_xlsx(self, template_processor, output_dir):
        """Test that all variables are replaced in XLSX."""
        from openpyxl import Workbook
        
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##var1##'
        ws['B1'] = '##var2##'
        ws['C1'] = '##var3##'
        ws['A2'] = 'Static: ##var1##'
        wb.save(str(template_path))
        
        output_path = output_dir / "output.xlsx"
        data = {
            'var1': 'ValueA',
            'var2': 'ValueB',
            'var3': 'ValueC'
        }
        
        template_processor.process_template(str(template_path), data, str(output_path))
        
        # Check all variables replaced
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        # Verify values
        assert result_ws['A1'].value == 'ValueA', "A1 not replaced"
        assert result_ws['B1'].value == 'ValueB', "B1 not replaced"
        assert result_ws['C1'].value == 'ValueC', "C1 not replaced"
        assert result_ws['A2'].value == 'Static: ValueA', "A2 not replaced"
        
        # Check for any remaining markers
        for row in result_ws.iter_rows(values_only=True):
            for cell in row:
                if cell and isinstance(cell, str):
                    assert '##' not in cell, f"Variable marker found: {cell}"
    
    def test_partial_substitution(self, template_processor, output_dir):
        """Test behavior when some variables are missing."""
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##provided## and ##missing##')
        doc.save(str(template_path))
        
        output_path = output_dir / "output.docx"
        data = {'provided': 'Value'}
        # Note: ##missing## not in data
        
        template_processor.process_template(str(template_path), data, str(output_path))
        
        result_doc = Document(str(output_path))
        text = '\n'.join([p.text for p in result_doc.paragraphs])
        
        # Provided should be replaced
        assert 'Value' in text
        # Missing should remain as variable
        assert '##missing##' in text


class TestExcelAutoAdjust:
    """Tests for Excel auto-adjust feature validation."""
    
    def test_auto_adjust_applied(self, template_processor, output_dir):
        """Test that auto-adjust is applied correctly."""
        from openpyxl import Workbook
        
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '##short##'
        ws['A2'] = '##long##'
        ws.column_dimensions['A'].width = 10
        wb.save(str(template_path))
        
        output_path = output_dir / "output.xlsx"
        data = {
            'short': 'Hi',
            'long': 'This is a very long text that should trigger width adjustment'
        }
        
        auto_adjust_options = {
            'auto_adjust_height': True,
            'auto_adjust_width': True
        }
        
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path),
            auto_adjust_options=auto_adjust_options
        )
        
        # Verify adjustments
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        # Width should be increased
        assert result_ws.column_dimensions['A'].width > 10, "Width not adjusted"
        
        # Height for row 2 should be auto (None)
        assert result_ws.row_dimensions[2].height is None, "Height not set to auto"
    
    def test_auto_adjust_range_only(self, template_processor, output_dir):
        """Test auto-adjust with specific range."""
        from openpyxl import Workbook
        
        template_path = output_dir / "template.xlsx"
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Outside'
        ws['B2'] = '##inside##'
        ws['C3'] = '##also_inside##'
        ws['D5'] = 'Also outside'
        wb.save(str(template_path))
        
        output_path = output_dir / "output.xlsx"
        data = {
            'inside': 'Inside range',
            'also_inside': 'Also inside range'
        }
        
        auto_adjust_options = {
            'auto_adjust_height': True,
            'auto_adjust_width': True,
            'adjust_range': 'B2:C3'
        }
        
        template_processor.process_template(
            str(template_path),
            data,
            str(output_path),
            auto_adjust_options=auto_adjust_options
        )
        
        result_wb = load_workbook(str(output_path))
        result_ws = result_wb.active
        
        # Verify data substituted
        assert result_ws['B2'].value == 'Inside range'
        assert result_ws['C3'].value == 'Also inside range'
        
        # Columns B and C should be adjusted
        assert result_ws.column_dimensions['B'].width is not None
        assert result_ws.column_dimensions['C'].width is not None


class TestZipArchives:
    """Tests for ZIP archive creation and integrity."""
    
    def test_zip_contains_all_files(self, job_manager, output_dir):
        """Test that ZIP contains all generated files."""
        from docx import Document
        from openpyxl import Workbook
        
        # Create template and data
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('##name##')
        doc.save(str(template_path))
        
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['name', 'filename'])
        ws.append(['Person 1', 'file1'])
        ws.append(['Person 2', 'file2'])
        ws.append(['Person 3', 'file3'])
        wb.save(str(data_path))
        
        # Process job
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx']
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Verify ZIP
        assert processed_job.zip_file_path is not None
        zip_path = Path(processed_job.zip_file_path)
        assert zip_path.exists(), "ZIP not created"
        
        # Check ZIP contents
        with zipfile.ZipFile(str(zip_path), 'r') as zf:
            files = zf.namelist()
            assert len(files) == 3, f"Expected 3 files, got {len(files)}"
            
            # All files should be DOCX
            for file in files:
                assert file.endswith('.docx'), f"Wrong file type: {file}"
    
    def test_zip_file_integrity(self, job_manager, output_dir):
        """Test ZIP file can be extracted and files are valid."""
        from docx import Document
        from openpyxl import Workbook
        
        template_path = output_dir / "template.docx"
        doc = Document()
        doc.add_paragraph('Test content ##value##')
        doc.save(str(template_path))
        
        data_path = output_dir / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(['value', 'filename'])
        ws.append(['ABC', 'test'])
        wb.save(str(data_path))
        
        job = job_manager.create_job(
            template_path=str(template_path),
            data_path=str(data_path),
            output_formats=['docx']
        )
        
        processed_job = job_manager.process_job(job.id)
        
        # Extract and validate
        zip_path = Path(processed_job.zip_file_path)
        extract_dir = output_dir / "extracted"
        extract_dir.mkdir(exist_ok=True)
        
        with zipfile.ZipFile(str(zip_path), 'r') as zf:
            zf.extractall(str(extract_dir))
        
        # Verify extracted files (use rglob to handle subdirectories)
        extracted_files = list(extract_dir.rglob('*.docx'))
        assert len(extracted_files) > 0, "No files extracted"
        
        # Check file is valid
        for file in extracted_files:
            doc = Document(str(file))
            text = '\n'.join([p.text for p in doc.paragraphs])
            assert 'ABC' in text, "Content not found in extracted file"


if __name__ == '__main__':
    pytest.main([__file__, '-v', '--tb=short'])
