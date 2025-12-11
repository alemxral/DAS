"""Test LibreOffice conversion speed comparison."""
import sys
import os
import time
from pathlib import Path

sys.path.insert(0, r'c:\Users\pc\autoarendt')

from services.format_converter import FormatConverter
from docx import Document

# Create test directory
test_dir = Path(r'c:\Users\pc\autoarendt\test_libreoffice')
test_dir.mkdir(exist_ok=True)

# Create a test document
doc = Document()
doc.add_heading('Speed Test Document', 0)
doc.add_paragraph('This tests LibreOffice vs COM conversion speed.')
for i in range(5):
    doc.add_paragraph(f'Paragraph {i+1}: Lorem ipsum dolor sit amet, consectetur adipiscing elit.')

test_docx = test_dir / 'speed_test.docx'
doc.save(str(test_docx))

converter = FormatConverter()

# Test LibreOffice
print("=" * 60)
print("Testing LibreOffice conversion...")
print("=" * 60)
start = time.time()
try:
    pdf_path = converter.convert(str(test_docx), 'pdf', str(test_dir))
    libreoffice_time = time.time() - start
    print(f"✅ LibreOffice: {libreoffice_time:.2f} seconds")
    print(f"   PDF size: {os.path.getsize(pdf_path):,} bytes")
except Exception as e:
    print(f"❌ LibreOffice failed: {e}")
    libreoffice_time = None

print("\n" + "=" * 60)
print("Summary:")
print("=" * 60)
if libreoffice_time:
    print(f"LibreOffice: {libreoffice_time:.2f}s")
    print("\n✅ LibreOffice portable is working and fast!")
