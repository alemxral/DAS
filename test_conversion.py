"""Test LibreOffice PDF conversion."""
import os
import sys
import time
from pathlib import Path

# Add project to path
sys.path.insert(0, r'c:\Users\pc\autoarendt')

from services.format_converter import FormatConverter

# Create a simple test document
test_dir = r'c:\Users\pc\autoarendt\test_output'
os.makedirs(test_dir, exist_ok=True)

# Create a simple Word document
from docx import Document
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test document for LibreOffice conversion.')
doc.add_paragraph('If you can read this in a PDF, the conversion worked!')

test_docx = os.path.join(test_dir, 'test.docx')
doc.save(test_docx)
print(f'Created test document: {test_docx}')

# Test conversion
print('\n--- Testing PDF Conversion ---')
converter = FormatConverter()

try:
    start = time.time()
    pdf_path = converter.convert(test_docx, 'pdf', test_dir)
    elapsed = time.time() - start
    
    print(f'\n[OK] Conversion successful!')
    print(f'PDF created: {pdf_path}')
    print(f'Time taken: {elapsed:.2f} seconds')
    print(f'PDF size: {os.path.getsize(pdf_path) / 1024:.1f} KB')
    
    if os.path.exists(pdf_path):
        print(f'\n[OK] PDF file exists and is ready!')
    else:
        print(f'\n[ERROR] PDF file not found!')
        
except Exception as e:
    print(f'\n[ERROR] Conversion failed: {e}')
    import traceback
    traceback.print_exc()

print(f'\nTest files in: {test_dir}')
